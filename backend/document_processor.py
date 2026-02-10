"""
Document Processing Module
Handles: parsing (PDF, DOCX, TXT), intelligent chunking, embedding, and vector storage.
"""

import os
import re
import uuid
import hashlib
from typing import List, Tuple

import chromadb
from sentence_transformers import SentenceTransformer
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


# ── Globals ─────────────────────────────────────────────────────────

_embed_model: SentenceTransformer = None
_chroma_client: chromadb.ClientAPI = None

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
CHROMA_DIR = os.path.join(os.path.dirname(__file__), "chroma_db")


def get_embed_model(model_name: str = "all-MiniLM-L6-v2") -> SentenceTransformer:
    global _embed_model
    if _embed_model is None:
        print(f"[DocProcessor] Loading embedding model: {model_name}")
        _embed_model = SentenceTransformer(model_name)
    return _embed_model


def get_chroma_client() -> chromadb.ClientAPI:
    global _chroma_client
    if _chroma_client is None:
        os.makedirs(CHROMA_DIR, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)
    return _chroma_client


# ── Text Extraction ─────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── Intelligent Chunking ────────────────────────────────────────────

# Logistics section headers commonly found in documents
SECTION_PATTERNS = [
    r"(?i)^(shipment\s+(?:details|information|summary))",
    r"(?i)^(shipper\s+(?:information|details|name))",
    r"(?i)^(consignee\s+(?:information|details|name))",
    r"(?i)^(carrier\s+(?:information|details|name))",
    r"(?i)^(rate\s+(?:confirmation|details|summary))",
    r"(?i)^(pickup\s+(?:details|information|date|time))",
    r"(?i)^(delivery\s+(?:details|information|date|time))",
    r"(?i)^(bill\s+of\s+lading)",
    r"(?i)^(freight\s+(?:charges|details|bill))",
    r"(?i)^(special\s+(?:instructions|notes|requirements))",
    r"(?i)^(equipment\s+(?:type|details|requirements))",
    r"(?i)^(terms\s+and\s+conditions)",
    r"(?i)^(payment\s+(?:terms|details))",
    r"(?i)^(insurance|claims|liability)",
    r"(?i)^(commodity|cargo|goods)\s",
    r"(?i)^(origin|destination)\s",
    r"(?i)^(weight|dimensions)\s",
    r"(?i)^#{1,3}\s",  # Markdown headers
]


def _is_section_boundary(line: str) -> bool:
    """Check if a line looks like a section header."""
    line_stripped = line.strip()
    if not line_stripped:
        return False
    for pattern in SECTION_PATTERNS:
        if re.match(pattern, line_stripped):
            return True
    # Heuristic: ALL CAPS lines that are short are likely headers
    if line_stripped.isupper() and 3 < len(line_stripped) < 80:
        return True
    return False


def intelligent_chunk(text: str, chunk_size: int = 512, chunk_overlap: int = 64) -> List[str]:
    """
    Chunk text intelligently:
    1. First splits by section boundaries (logistics headers).
    2. Then splits large sections by sentence boundaries.
    3. Finally falls back to character-level splitting with overlap.
    """
    if not text.strip():
        return []

    # Step 1: Split into sections by headers
    lines = text.split("\n")
    sections: List[str] = []
    current_section: List[str] = []

    for line in lines:
        if _is_section_boundary(line) and current_section:
            sections.append("\n".join(current_section))
            current_section = [line]
        else:
            current_section.append(line)

    if current_section:
        sections.append("\n".join(current_section))

    # Step 2: Break sections into chunks
    chunks: List[str] = []

    for section in sections:
        section = section.strip()
        if not section:
            continue

        if len(section) <= chunk_size:
            chunks.append(section)
        else:
            # Split by sentences first
            sentences = re.split(r'(?<=[.!?])\s+', section)
            current_chunk = ""

            for sentence in sentences:
                if len(current_chunk) + len(sentence) + 1 <= chunk_size:
                    current_chunk += (" " if current_chunk else "") + sentence
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    # If single sentence is longer than chunk_size, split by chars
                    if len(sentence) > chunk_size:
                        for i in range(0, len(sentence), chunk_size - chunk_overlap):
                            sub = sentence[i:i + chunk_size]
                            if sub.strip():
                                chunks.append(sub.strip())
                        current_chunk = ""
                    else:
                        current_chunk = sentence

            if current_chunk.strip():
                chunks.append(current_chunk.strip())

    # Filter out very small chunks (less than 20 chars)
    chunks = [c for c in chunks if len(c) >= 20]

    return chunks


# ── Embedding & Storage ─────────────────────────────────────────────

def process_and_store(
    file_path: str,
    filename: str,
    chunk_size: int = 512,
    chunk_overlap: int = 64,
    embedding_model: str = "all-MiniLM-L6-v2",
) -> Tuple[str, int]:
    """
    Full pipeline: extract → chunk → embed → store in ChromaDB.
    Returns (document_id, num_chunks).
    """
    # Generate a unique document ID
    doc_id = hashlib.md5(f"{filename}-{uuid.uuid4()}".encode()).hexdigest()[:16]

    # Extract text
    raw_text = extract_text(file_path)
    if not raw_text.strip():
        raise ValueError("No text could be extracted from the document.")

    # Chunk
    chunks = intelligent_chunk(raw_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    if not chunks:
        raise ValueError("Document produced no valid chunks after processing.")

    # Embed
    model = get_embed_model(embedding_model)
    embeddings = model.encode(chunks, show_progress_bar=False, normalize_embeddings=True)

    # Store in ChromaDB
    client = get_chroma_client()
    collection = client.get_or_create_collection(
        name=f"doc_{doc_id}",
        metadata={"hnsw:space": "cosine"},
    )

    ids = [f"chunk_{i}" for i in range(len(chunks))]
    metadatas = [{"chunk_index": i, "filename": filename, "doc_id": doc_id} for i in range(len(chunks))]

    collection.add(
        documents=chunks,
        embeddings=embeddings.tolist(),
        ids=ids,
        metadatas=metadatas,
    )

    # Also store the raw text for structured extraction
    raw_path = os.path.join(UPLOAD_DIR, f"{doc_id}_raw.txt")
    with open(raw_path, "w", encoding="utf-8") as f:
        f.write(raw_text)

    print(f"[DocProcessor] Stored {len(chunks)} chunks for document {doc_id}")
    return doc_id, len(chunks)


def retrieve_chunks(doc_id: str, query: str, top_k: int = 5, embedding_model: str = "all-MiniLM-L6-v2") -> List[dict]:
    """
    Retrieve the top-k most relevant chunks for a given query.
    Returns list of {text, chunk_index, similarity_score}.
    """
    model = get_embed_model(embedding_model)
    query_embedding = model.encode([query], normalize_embeddings=True).tolist()

    client = get_chroma_client()
    collection = client.get_collection(name=f"doc_{doc_id}")

    results = collection.query(
        query_embeddings=query_embedding,
        n_results=min(top_k, collection.count()),
        include=["documents", "distances", "metadatas"],
    )

    chunks = []
    if results and results["documents"] and results["documents"][0]:
        for i, (doc, dist, meta) in enumerate(
            zip(results["documents"][0], results["distances"][0], results["metadatas"][0])
        ):
            # ChromaDB returns cosine distance; convert to similarity
            similarity = 1.0 - dist
            chunks.append({
                "text": doc,
                "chunk_index": meta.get("chunk_index", i),
                "similarity_score": round(similarity, 4),
            })

    return chunks


def get_raw_text(doc_id: str) -> str:
    """Load the full raw text of a document for structured extraction."""
    raw_path = os.path.join(UPLOAD_DIR, f"{doc_id}_raw.txt")
    if not os.path.exists(raw_path):
        raise FileNotFoundError(f"Raw text not found for document {doc_id}")
    with open(raw_path, "r", encoding="utf-8") as f:
        return f.read()


def document_exists(doc_id: str) -> bool:
    """Check if a document exists in the vector store."""
    try:
        client = get_chroma_client()
        client.get_collection(name=f"doc_{doc_id}")
        return True
    except Exception:
        return False
